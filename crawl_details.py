import json, time, random, datetime, re, sys, traceback, os
import threading
from queue import Queue, Empty

import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ====== Cấu hình ======
MAX_WORKERS = 5          # số luồng (số Chrome chạy song song). Tùy máy, 2-5 là hợp lý
PAGELOAD_TIMEOUT = 45    # giây chờ trang
WAIT_TABLE_TIMEOUT = 25  # giây chờ bảng chi tiết
RETRY_PER_ITEM = 2       # số lần retry cho mỗi công ty khi lỗi tạm thời
HEADLESS = False         # True để bật chế độ ẩn chrome
OUTFILE_PREFIX = "Vu"    # tiền tố tên file docx xuất ra
# ======================

DETAIL_FIELDS = {
    "Số ĐKKD/MST": "Số ĐKKD/MST",
    "Ngày hoạt động": "Ngày hoạt động",
    "Tình trạng": "Tình trạng",
    "Địa chỉ": "Địa chỉ",
    "Người đại diện": "Người đại diện",
    "Điện thoại": "Điện thoại",
}

# Các đầu số hợp lệ
VALID_PREFIXES = ["090", "093", "089", "070"]
VALID_PREFIX_RANGES = [(32, 39), (76, 79)]  # 032..039 và 076..079

def is_valid_phone(phone: str) -> bool:
    """Lọc số ĐT theo đầu số cho phép."""
    if not phone:
        return False
    digits = re.sub(r"\D", "", phone)
    if len(digits) < 3:
        return False

    prefix3 = digits[:3]
    if prefix3 in VALID_PREFIXES:
        return True

    try:
        prefix_int = int(prefix3)
        for low, high in VALID_PREFIX_RANGES:
            if low <= prefix_int <= high:
                return True
    except:
        return False

    return False

driver_lock = threading.Lock()

def build_driver():
    """Khởi tạo UC Chrome cho từng luồng."""
    options = uc.ChromeOptions()
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("--disable-infobars")
    options.add_argument("--start-maximized")
    options.add_argument(
        "user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115 Safari/537.36"
    )
    if HEADLESS:
        options.add_argument("--headless=new")
        options.add_argument("--disable-gpu")
        options.add_argument("--window-size=1920,1080")

    with driver_lock:
        driver = uc.Chrome(options=options, use_subprocess=True)
    return driver

def gentle_scroll(driver):
    """Cuộn nhẹ để kích hoạt lazy-load."""
    try:
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight/2);")
        time.sleep(random.uniform(0.8, 1.6))
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(random.uniform(0.8, 1.6))
        driver.execute_script("window.scrollTo(0, 0);")
    except Exception:
        pass

def cloudflare_guard(driver):
    """Phát hiện & xử lý khi bị Cloudflare challenge."""
    try:
        html = driver.page_source or ""
        if ("Checking your browser" in html) or ("cf-browser-verification" in html):
            print("⚠️  Cloudflare chặn, đợi xử lý...")
            time.sleep(random.uniform(8, 12))
            driver.refresh()
            time.sleep(random.uniform(3, 6))
            return True
    except Exception:
        pass
    return False

def parse_details(driver):
    details = {field: None for field in DETAIL_FIELDS}

    try:
        tax_id = driver.find_element(By.CSS_SELECTOR, "td[itemprop='taxID']").text.strip()
        if tax_id:
            details["Số ĐKKD/MST"] = tax_id
    except:
        pass

    try:
        addr = driver.find_element(
            By.CSS_SELECTOR, "table.company-table td[itemprop='address']"
        ).text.strip()
        if addr:
            details["Địa chỉ"] = addr
    except:
        pass

    owner = None
    try:
        owner = driver.find_element(
            By.CSS_SELECTOR,
            "table.company-table tr[itemprop='Owner'] td:nth-child(2) a"
        ).text.strip()
    except:
        try:
            owner = driver.find_element(
                By.CSS_SELECTOR,
                "table.company-table tr[itemprop='Owner'] td:nth-child(2) span[itemprop='Owner']"
            ).text.strip()
        except:
            pass
    if owner:
        details["Người đại diện"] = owner

    for row in driver.find_elements(By.CSS_SELECTOR, "table.company-table > tbody > tr"):
        tds = row.find_elements(By.TAG_NAME, "td")
        if len(tds) != 2:
            continue
        key = tds[0].text.strip().replace(":", "")
        val = tds[1].text.strip()

        if key in ("Ngày cấp", "Ngày hoạt động", "Tình trạng", "Điện thoại"):
            val = val.split("\n")[0].strip()
            if val:
                details[key] = val

    return details

def get_company_details(driver, url):
    driver.get(url)
    time.sleep(random.uniform(1.5, 3.0))
    gentle_scroll(driver)
    cloudflare_guard(driver)

    try:
        WebDriverWait(driver, WAIT_TABLE_TIMEOUT).until(
            EC.presence_of_element_located((
                By.CSS_SELECTOR,
                "table.company-table td[itemprop='address'], "
                "table.company-table tr[itemprop='Owner']"
            ))
        )
    except Exception:
        print("❌ Không thấy bảng chi tiết.")
        return {field: None for field in DETAIL_FIELDS}

    return parse_details(driver)

def worker(worker_id: int, q: Queue, results_list: list, results_lock: threading.Lock):
    """Luồng công nhân."""
    driver = None
    try:
        driver = build_driver()
        while True:
            try:
                item = q.get(timeout=2)
            except Empty:
                break

            idx, company = item
            name = company.get("name", "").strip()
            link = company.get("link", "").strip()
            if not link:
                print(f"[W{worker_id}] ⚠️ Bỏ qua item không có link: {name}")
                q.task_done()
                continue

            print(f"[W{worker_id}] ▶️  ({idx}) Đang lấy: {name}")

            success = False
            last_err = None
            for attempt in range(1, RETRY_PER_ITEM + 1):
                try:
                    details = get_company_details(driver, link)
                    phone = details.get("Điện thoại")
                    status = (details.get("Tình trạng") or "").strip().lower()

                    # chỉ nhận "đang hoạt động"
                    if status != "đang hoạt động":
                        print(f"[W{worker_id}] ⛔ Bỏ qua {name} (Tình trạng: {status})")
                    elif not is_valid_phone(phone):
                        print(f"[W{worker_id}] ⛔ Bỏ qua {name} (SĐT không hợp lệ: {phone})")
                    else:
                        merged = dict(company)
                        merged.update(details)
                        with results_lock:
                            results_list.append(merged)
                        print(f"[W{worker_id}] ✅ OK: {name}")
                    success = True
                    break
                except Exception as e:
                    last_err = e
                    print(f"[W{worker_id}] ❌ Lỗi {name} (lần {attempt}/{RETRY_PER_ITEM}): {e}")
                    traceback.print_exc(limit=1)
                    try:
                        driver.refresh()
                    except Exception:
                        pass
                    time.sleep(random.uniform(2, 4))

            if not success and last_err:
                print(f"[W{worker_id}] 🚫 Bỏ qua {name} sau khi retry: {last_err}")

            time.sleep(random.uniform(1.5, 3.5))
            q.task_done()

    except Exception as e:
        print(f"[W{worker_id}] 💥 Lỗi luồng: {e}")
        traceback.print_exc(limit=1)
    finally:
        if driver:
            try:
                driver.quit()
            except Exception:
                pass
        print(f"[W{worker_id}] 🔚 Đã đóng driver.")

def export_to_word(items: list, outfile_path: str):
    """Xuất kết quả ra Word với định dạng chuẩn."""
    doc = Document()

    # Font mặc định
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(9)

    p_format = style.paragraph_format
    p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_format.space_after = Pt(0)

    # Nội dung
    for comp in items:
        name = comp.get("name", "")
        p = doc.add_paragraph()
        run = p.add_run((name or "").upper())
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for field in DETAIL_FIELDS.keys():
            value = comp.get(field, "")
            if value:
                doc.add_paragraph(f"{field}: {value}")
        doc.add_paragraph("")
    doc.save(outfile_path)

def normalize_phone(phone: str) -> str:
    """Chuẩn hóa số điện thoại: giữ lại chỉ chữ số."""
    if not phone:
        return ""
    return re.sub(r"\D", "", phone)

def generate_word_filename(prefix="Vu"):
    today = datetime.datetime.now().strftime("%d.%m")
    i = 1
    while True:
        filename = f"{prefix}.{today}.{i}.docx"
        if not os.path.exists(filename):
            return filename
        i += 1

def main():
    try:
        with open("companies.json", "r", encoding="utf-8") as f:
            companies = json.load(f)
    except Exception as e:
        print(f"❌ Không đọc được companies.json: {e}")
        sys.exit(1)

    if not isinstance(companies, list) or not companies:
        print("⚠️ Danh sách companies rỗng hoặc sai định dạng.")
        sys.exit(0)

    q = Queue()
    for i, comp in enumerate(companies, start=1):
        q.put((i, comp))

    results = []
    results_lock = threading.Lock()

    workers = []
    n_workers = max(1, min(MAX_WORKERS, q.qsize()))
    print(f"🚀 Khởi động {n_workers} luồng ...")
    for wid in range(1, n_workers + 1):
        t = threading.Thread(target=worker, args=(wid, q, results, results_lock), daemon=True)
        t.start()
        workers.append(t)

    for t in workers:
        t.join()

    # loại trùng
    seen_links = set()
    deduped = []
    for item in results:
        link = item.get("link")
        name = item.get("name")

        key = link or name
        if key and key in seen_links:
            continue
        if key:
            seen_links.add(key)
        deduped.append(item)

    outfile = generate_word_filename(OUTFILE_PREFIX)
    export_to_word(deduped, outfile)

    print(f"✅ Hoàn tất! Tổng hợp {len(deduped)}/{len(companies)} mục hợp lệ.")
    print(f"📝 Đã lưu kết quả vào {outfile}")

if __name__ == "__main__":
    main()
